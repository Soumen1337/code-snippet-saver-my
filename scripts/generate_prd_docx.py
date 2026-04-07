"""
Generate a professional Word document for the SnippetVault PRD.
Uses python-docx with custom styling for a polished, presentation-ready output.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colors ──────────────────────────────────────────────
VIOLET = RGBColor(0x7C, 0x3A, 0xED)
CYAN = RGBColor(0x06, 0xB6, 0xD4)
DARK_BG = RGBColor(0x0A, 0x0A, 0x0F)
DARK_TEXT = RGBColor(0x1D, 0x1D, 0x1F)
GRAY_TEXT = RGBColor(0x6E, 0x6E, 0x73)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF5, 0xF5, 0xF7)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
TABLE_HEADER_BG = "7C3AED"
TABLE_ALT_BG = "F5F3FF"
BORDER_COLOR = "E5E7EB"

doc = Document()

# ── Page Setup ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Custom Styles ───────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading 1
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(26)
h1_style.font.bold = True
h1_style.font.color.rgb = VIOLET
h1_style.paragraph_format.space_before = Pt(24)
h1_style.paragraph_format.space_after = Pt(12)

# Heading 2
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(18)
h2_style.font.bold = True
h2_style.font.color.rgb = DARK_TEXT
h2_style.paragraph_format.space_before = Pt(20)
h2_style.paragraph_format.space_after = Pt(8)

# Heading 3
h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Calibri'
h3_style.font.size = Pt(14)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)  # violet-800
h3_style.paragraph_format.space_before = Pt(16)
h3_style.paragraph_format.space_after = Pt(6)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = DARK_TEXT
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.line_spacing = 1.0

# Subtitle
subtitle_style = doc.styles.add_style('Subtitle2', WD_STYLE_TYPE.PARAGRAPH)
subtitle_style.font.name = 'Calibri'
subtitle_style.font.size = Pt(12)
subtitle_style.font.color.rgb = GRAY_TEXT
subtitle_style.font.italic = True

# Quote
quote_style = doc.styles.add_style('Quote2', WD_STYLE_TYPE.PARAGRAPH)
quote_style.font.name = 'Calibri'
quote_style.font.size = Pt(11)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
quote_style.paragraph_format.left_indent = Cm(1)
quote_style.paragraph_format.space_before = Pt(8)
quote_style.paragraph_format.space_after = Pt(8)

# ── Helper Functions ────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color="D4D4D8"):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def add_styled_table(headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER_BG)
        set_cell_borders(cell, TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            set_cell_borders(cell, BORDER_COLOR)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    doc.add_paragraph("")  # spacing after table
    return table

def add_bullet(text, level=0, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)

def add_body(text):
    """Add body paragraph."""
    p = doc.add_paragraph(text)
    return p

def add_divider():
    """Add a visual divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="E5E7EB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_note_box(text, prefix="Note"):
    """Add a highlighted note/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"⚠️ {prefix}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY_TEXT

def add_code_block(lines):
    """Add a code block with background styling."""
    for line in lines:
        p = doc.add_paragraph(line, style='CodeBlock')
        # Add gray background
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F5" w:val="clear"/>')
        pPr.append(shading)


# ═══════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════

# ── COVER PAGE ──────────────────────────────────────────

# Add spacing before title
for _ in range(6):
    doc.add_paragraph("")

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("SnippetVault")
run.font.size = Pt(42)
run.bold = True
run.font.color.rgb = VIOLET
run.font.name = 'Calibri'

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Product Requirements Document")
run.font.size = Pt(20)
run.font.color.rgb = GRAY_TEXT
run.font.name = 'Calibri'

# Tagline
tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_p.paragraph_format.space_before = Pt(16)
run = tag_p.add_run("Save, organize, and version-control your code snippets")
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = CYAN
run.font.name = 'Calibri'

# Metadata
doc.add_paragraph("")
doc.add_paragraph("")
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_before = Pt(30)
for label, value in [
    ("Version: ", "1.0"),
    ("  |  Date: ", "April 7, 2026"),
    ("  |  Status: ", "Prototype — Phase 1"),
]:
    run = meta_p.add_run(label)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    run = meta_p.add_run(value)
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = DARK_TEXT

# Cost badge
cost_p = doc.add_paragraph()
cost_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cost_p.paragraph_format.space_before = Pt(20)
run = cost_p.add_run("💵 Total Infrastructure Cost: $0")
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = GREEN

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Built entirely on free-tier services — Supabase Free Plan + Vercel Hobby Plan")
run.font.size = Pt(10)
run.font.color.rgb = GRAY_TEXT

# Page break
doc.add_page_break()


# ── TABLE OF CONTENTS ───────────────────────────────────
doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1. Executive Summary",
    "2. Problem Statement",
    "3. Motivation & Vision",
    "4. Target Users",
    "5. Use Cases — How, When, Where",
    "6. Feature Specification",
    "7. User Flows",
    "8. Application Architecture & Flow",
    "9. Technology Stack",
    "10. Data Model & Schema Design",
    "11. API Design",
    "12. Security & Access Control",
    "13. Development Process & Timeline",
    "14. Future Roadmap (Vision Only)",
    "15. Success Metrics",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT

doc.add_page_break()


# ── 1. EXECUTIVE SUMMARY ───────────────────────────────
doc.add_heading("1. Executive Summary", level=1)

add_body(
    "SnippetVault is a modern, full-stack web application that allows developers to save, organize, search, "
    "and version-control their code snippets. Think of it as a personal Git repository specifically designed "
    "for reusable code fragments."
)

add_body("Unlike plain note-taking apps or GitHub Gists, SnippetVault provides:")

add_bullet("Git-like version history — every edit creates a tracked version with a commit message")
add_bullet("Side-by-side diff comparison — see exactly what changed between any two versions")
add_bullet("Smart tagging & multi-filter search — find any snippet in milliseconds")
add_bullet("Syntax highlighting — for 25+ programming languages")
add_bullet("One-click copy — paste into any IDE or terminal instantly")

add_body(
    "The application is built with Next.js 16, Supabase (PostgreSQL + Auth), and a modern component library. "
    "The entire stack runs on 100% free-tier infrastructure — Supabase Free Plan + Vercel Hobby Plan — "
    "with zero cost to build, deploy, and run."
)

add_divider()


# ── 2. PROBLEM STATEMENT ───────────────────────────────
doc.add_heading("2. Problem Statement", level=1)

doc.add_heading("The Core Problem", level=2)
add_body(
    "Developers constantly write, reuse, and evolve small pieces of code across projects — "
    "utility functions, database queries, regex patterns, CSS snippets, Docker configs, and more."
)

add_body("Where does this code live today?")

add_styled_table(
    ["Current Solution", "Problem"],
    [
        ["Scattered text files", "No search, no organization, easily lost"],
        ["GitHub Gists", "No version history with diffs, poor discoverability"],
        ["IDE snippets", "Locked to one editor, not portable"],
        ["Notion / Google Docs", "No syntax highlighting, no code-aware features"],
        ["Memory", "Unreliable — developers waste time rewriting solved problems"],
    ],
    col_widths=[2.5, 4.0]
)

doc.add_heading("Impact", level=2)
add_body(
    "A 2023 Stack Overflow Developer Survey found that developers spend ~30% of their time searching for "
    "or rewriting code they've written before. For a team of 10 developers, this translates to an estimated "
    "3 full-time engineers' worth of productivity lost per year."
)

doc.add_heading("The Gap", level=2)
add_body(
    "No existing tool combines code-specific features (syntax highlighting, version history, diff viewer) "
    "with personal knowledge management (tags, search, organization) in a lightweight, browser-based interface."
)

add_divider()


# ── 3. MOTIVATION & VISION ─────────────────────────────
doc.add_heading("3. Motivation & Vision", level=1)

p = doc.add_paragraph(style='Quote2')
p.add_run(
    '"Every developer has a personal library of code patterns. '
    'SnippetVault makes that library searchable, versioned, and always accessible."'
)

doc.add_heading("Motivation", level=2)

p = doc.add_paragraph()
run = p.add_run("1. Academic Context — ")
run.bold = True
p.add_run(
    "This project serves as a prototype for a Software Engineering course, demonstrating full-stack "
    "development skills including authentication, CRUD operations, real-time search, and version control system design."
)

p = doc.add_paragraph()
run = p.add_run("2. Real Developer Pain — ")
run.bold = True
p.add_run(
    "The problem is genuine. The team members themselves have experienced the frustration of losing "
    "useful code snippets or rewriting utility functions from scratch."
)

p = doc.add_paragraph()
run = p.add_run("3. Technical Learning — ")
run.bold = True
p.add_run("The project exercises advanced concepts:")

add_bullet("Server-side rendering with Next.js App Router", level=1)
add_bullet("Row-level security with Supabase/PostgreSQL", level=1)
add_bullet("Optimistic UI updates with SWR", level=1)
add_bullet("Diff algorithms implemented in JavaScript", level=1)
add_bullet("Responsive, accessible UI with Radix primitives", level=1)

doc.add_heading("Vision", level=2)

add_styled_table(
    ["Phase", "Scope", "Cost", "Status"],
    [
        ["Phase 1 (Current)", "Personal snippet manager with version control", "$0 — Free tier", "✅ Complete"],
        ["Phase 2 (Future)", "AI-powered auto-tagging, natural language search", "💰 Requires paid APIs", "📋 Vision only"],
        ["Phase 3 (Future)", "Team collaboration, shared collections, marketplace", "💰 Requires Pro plan", "📋 Vision only"],
    ],
    col_widths=[1.5, 2.5, 1.5, 1.0]
)

add_note_box(
    "Phases 2 and 3 are documented for academic discussion and growth potential only. "
    "They are NOT part of the current prototype and would require paid services.",
    "Important"
)

add_divider()


# ── 4. TARGET USERS ─────────────────────────────────────
doc.add_heading("4. Target Users", level=1)

doc.add_heading("Primary Personas", level=2)

doc.add_heading("👨‍💻 The Student Developer", level=3)
add_bullet("18–24 years old, learning multiple languages simultaneously", bold_prefix="Profile: ")
add_bullet("Store solutions from assignments, tutorials, and side projects", bold_prefix="Needs: ")
add_bullet("Code from last semester is gone — lost in old folders or deleted repos", bold_prefix="Pain Point: ")

doc.add_heading("👩‍💻 The Professional Developer", level=3)
add_bullet("Software engineer (25–40), works across multiple projects and clients", bold_prefix="Profile: ")
add_bullet("Maintain a personal utility library, track how solutions evolve", bold_prefix="Needs: ")
add_bullet('"I wrote a perfect date parser last month — where is it?"', bold_prefix="Pain Point: ")

doc.add_heading("🧑‍🏫 The Technical Writer / Educator", level=3)
add_bullet("Creates tutorials, documentation, or teaches coding workshops", bold_prefix="Profile: ")
add_bullet("Maintain a curated library of example code in multiple languages", bold_prefix="Needs: ")
add_bullet("Example code drifts out of date; no way to track which version was used in which tutorial", bold_prefix="Pain Point: ")

doc.add_heading("User Demographics", level=2)
add_styled_table(
    ["Attribute", "Value"],
    [
        ["Technical Level", "Intermediate to advanced developers"],
        ["Primary Device", "Desktop / Laptop (code-centric workflow)"],
        ["Browser", "Chrome, Firefox, Edge, Safari"],
        ["Frequency of Use", "Daily (during active development)"],
        ["Session Length", "2–15 minutes per session"],
    ],
    col_widths=[2.5, 4.0]
)

add_divider()


# ── 5. USE CASES ────────────────────────────────────────
doc.add_heading("5. Use Cases — How, When, Where", level=1)

doc.add_heading("When Will Users Use SnippetVault?", level=2)
add_styled_table(
    ["Scenario", "Trigger", "Action"],
    [
        ["Learning a new concept", "User finds a useful pattern in a tutorial", "Save snippet → tag with topic and language"],
        ["Solving a bug", "User writes a clever fix", "Save the fix → add commit message explaining reasoning"],
        ["Starting a new project", "User needs boilerplate code", "Search snippets by language/tag → copy to project"],
        ["Refactoring existing code", "User improves a utility function", "Update snippet → version history tracks the change"],
        ["Code review prep", "User wants to compare approaches", "Open diff viewer → compare v1 vs v3 side-by-side"],
        ["Switching devices", "User moves from office to home", "Open SnippetVault in browser → all snippets synced"],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

doc.add_heading("Where Will It Be Used?", level=2)
add_bullet("In the browser alongside an IDE (split screen workflow)")
add_bullet("On any device with a modern web browser (cloud-based, no local install)")
add_bullet("In classrooms as a teaching tool for code versioning concepts")
add_bullet("In hackathons for quick access to reusable starter code")

doc.add_heading("How Will It Be Used?", level=2)
steps = [
    ("Save: ", "Press 'New Snippet' → enter code, title, language, tags → save"),
    ("Find: ", "Type in search bar or click a language/tag filter → results appear instantly"),
    ("Use: ", "Click a snippet card → view full code → one-click copy to clipboard"),
    ("Evolve: ", "Edit a snippet → add a commit message → new version is tracked"),
    ("Compare: ", "Select two versions in history → view side-by-side diff"),
]
for bold, text in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(bold)
    run.bold = True
    run.font.color.rgb = VIOLET
    p.add_run(text)

add_divider()


# ── 6. FEATURE SPECIFICATION ───────────────────────────
doc.add_heading("6. Feature Specification", level=1)

doc.add_heading("6.1 Core Features (Implemented ✅)", level=2)

features = [
    ("F1: User Authentication", [
        "Email/password registration and login via Supabase Auth",
        "Secure session management with cookie-based tokens",
        "Protected routes via Next.js middleware",
        "Sign-up email confirmation flow",
        "Sign-out functionality",
    ]),
    ("F2: Snippet CRUD Operations", [
        "Create snippet with title, description, language, code content, and tags",
        "Read snippet list with preview (first 6 lines), or full detail in side panel",
        "Update snippet metadata and code — content changes trigger a new version",
        "Delete snippet with confirmation dialog (cascade deletes versions and tags)",
    ]),
    ("F3: Version Control System", [
        "Every content edit creates a new snippet_versions record",
        "Auto-incrementing version numbers (v1, v2, v3...)",
        "Optional commit messages (defaults to 'Version N')",
        "Full version history list in the snippet detail panel",
        "Ability to view any historical version's code",
    ]),
    ("F4: Side-by-Side Diff Viewer", [
        "Select any two versions from the history panel",
        "View additions (green) and deletions (red) in a unified diff format",
        "Built using the 'diff' library for line-by-line comparison",
    ]),
    ("F5: Smart Tagging System", [
        "User-created tags (each user has their own tag namespace)",
        "Many-to-many relationship (snippet ↔ tags)",
        "AND-logic tag filtering (selected tags must ALL match)",
        "Tag management API",
    ]),
    ("F6: Full-Text Search", [
        "Searches across title, description, and current_content",
        "Debounced input (300ms) to prevent excessive API calls",
        "Combined with language and tag filters for precise results",
    ]),
    ("F7: Language Filtering", [
        "Inline chip-based filter for quick language selection",
        "Sidebar navigation with 10 most common languages",
        "Full dropdown with 25+ supported languages",
    ]),
    ("F8: Syntax Highlighting", [
        "Powered by prism-react-renderer",
        "Supports 27 languages including JavaScript, TypeScript, Python, Rust, Go, and more",
    ]),
    ("F9: One-Click Copy", [
        "Clipboard copy button on every snippet card (appears on hover)",
        "Clipboard copy button in the detail panel",
        "Visual feedback with checkmark confirmation (2 seconds)",
    ]),
    ("F10: Dark/Light Mode", [
        "System-preference-aware theme detection",
        "Manual toggle in the sidebar",
        "Persistent theme selection across sessions",
    ]),
]

for title, items in features:
    doc.add_heading(title, level=3)
    for item in items:
        add_bullet(item)


doc.add_heading("6.2 Future Scope — Potential Enhancements", level=2)

add_note_box(
    "The features below are NOT part of the current prototype. They are documented purely for "
    "academic discussion — to demonstrate product thinking and growth potential. Many require paid APIs, "
    "significant development time, or infrastructure upgrades. None of these are planned for immediate implementation.",
    "Important Disclaimer"
)

add_styled_table(
    ["Feature", "Description", "Cost / Effort", "Why It's Future Scope"],
    [
        ["AI Auto-Tagging", "Analyze code → suggest tags", "💰 OpenAI API ($0.002–$0.06/call)", "Paid API, needs billing"],
        ["AI Natural Language Search", "Search using plain English", "💰 Embedding API + pgvector", "Paid API, complex setup"],
        ["Code Execution Sandbox", "Run JS/Python in-browser", "⏱️ 2–3 weeks dev time", "Heavy engineering effort"],
        ["Import from GitHub Gists", "OAuth + Gist import", "⏱️ 1 week + OAuth setup", "Moderate effort"],
        ["Export & Sharing", "Public links / embed widgets", "⏱️ 1 week", "Needs public DB rows"],
        ["Snippet Collections", "Folders for organizing", "⏱️ 1 week", "New DB table + UI"],
        ["Keyboard Shortcuts", "⌘K search, N new snippet", "🟢 Free, ~2 hours", "Low effort"],
        ["Analytics Dashboard", "Language charts, activity", "🟢 Free (recharts installed)", "Low effort"],
        ["Browser Extension", "Save from any webpage", "⏱️ 2–3 weeks", "Separate project"],
        ["VS Code Extension", "Save/search from IDE", "⏱️ 2–3 weeks", "Separate project"],
    ],
    col_widths=[1.5, 1.8, 1.8, 1.5]
)

p = doc.add_paragraph()
run = p.add_run("Bottom line: ")
run.bold = True
p.add_run("The current prototype delivers a complete, functional product using only free tools. "
          "The features above represent where the product could go — not where it needs to go right now.")

add_divider()


# ── 7. USER FLOWS ───────────────────────────────────────
doc.add_heading("7. User Flows", level=1)

doc.add_heading("7.1 New User Registration Flow", level=2)
flow_lines = [
    "Landing Page (/)",
    "  │",
    "  └── Click 'Get Started'",
    "       │",
    "       ▼",
    "Sign Up Page (/auth/sign-up)",
    "  │ Enter email + password → Submit",
    "  ▼",
    "Supabase sends confirmation email",
    "  ▼",
    "Sign Up Success Page (/auth/sign-up-success)",
    "  │ User checks email → clicks confirmation link",
    "  ▼",
    "Login Page (/auth/login)",
    "  │ Enter credentials",
    "  ▼",
    "Dashboard (/dashboard)",
    '  └── Empty state: "No snippets yet" + CTA',
]
add_code_block(flow_lines)

doc.add_heading("7.2 Create Snippet Flow", level=2)
flow_lines = [
    "Dashboard → Click 'New Snippet'",
    "  ▼",
    "Snippet Dialog (Modal)",
    "  │ Enter title, description, language",
    "  │ Write/paste code in editor",
    "  │ Select tags, enter commit message",
    "  │ Click 'Save'",
    "  ▼",
    "POST /api/snippets",
    "  ├── Creates snippets row",
    "  ├── Creates snippet_versions row (v1)",
    "  └── Creates snippet_tags junction rows",
    "  ▼",
    "Dashboard refreshes via SWR mutate()",
    "  └── New snippet card appears in grid",
]
add_code_block(flow_lines)

doc.add_heading("7.3 Edit & Version Flow", level=2)
flow_lines = [
    "Click snippet card → Detail Panel opens",
    "  ├── View current code (Code tab)",
    "  ├── View version history (History tab)",
    "  └── Click 'Edit'",
    "       ▼",
    "Snippet Dialog with pre-filled data",
    "  │ Modify code, enter commit message",
    "  │ Click 'Save'",
    "  ▼",
    "PUT /api/snippets/[id]",
    "  ├── Updates snippets.current_content",
    "  ├── Creates new snippet_versions (v2, v3...)",
    "  └── Updates tags (delete old, insert new)",
    "  ▼",
    "Dashboard + Detail Panel refresh",
]
add_code_block(flow_lines)

doc.add_heading("7.4 Search & Filter Flow", level=2)
flow_lines = [
    "Dashboard",
    "  ├── Type in search → debounced (300ms)",
    "  │   └── GET /api/snippets?search=<query>",
    "  ├── Click language chip (e.g., 'Python')",
    "  │   └── GET /api/snippets?language=python",
    "  ├── Select tag from filter",
    "  │   └── GET /api/snippets?tags=<id1>,<id2>",
    "  └── All filters combine:",
    "      └── GET /api/snippets?search=sort&language=js&tags=abc",
    "  ▼",
    "Grid updates with filtered results",
]
add_code_block(flow_lines)

doc.add_heading("7.5 Diff Comparison Flow", level=2)
flow_lines = [
    "Detail Panel → History tab",
    "  ├── Click Version 1 → selected",
    "  ├── Click Version 3 → selected",
    "  ▼",
    "'Diff' tab appears in tab bar → Click it",
    "  ▼",
    "Side-by-side diff view:",
    "  ├── Left:  Version 1 content",
    "  ├── Right: Version 3 content",
    "  ├── Green: Added lines",
    "  └── Red:   Removed lines",
]
add_code_block(flow_lines)

add_divider()


# ── 8. ARCHITECTURE ─────────────────────────────────────
doc.add_heading("8. Application Architecture & Flow", level=1)

doc.add_heading("8.1 High-Level Architecture", level=2)

add_styled_table(
    ["Layer", "Component", "Technology", "Responsibility"],
    [
        ["Client", "Landing Page", "Next.js (SSR)", "Marketing, SEO-optimized"],
        ["Client", "Auth Pages", "Next.js (SSR)", "Login, sign-up, error handling"],
        ["Client", "Dashboard", "React (CSR) + SWR", "Snippet CRUD, search, filters"],
        ["Server", "Middleware", "Next.js middleware", "Session refresh, route protection"],
        ["Server", "API Routes", "Next.js API handlers", "CRUD operations, auth checks"],
        ["Database", "PostgreSQL", "Supabase (Free Plan)", "Snippets, versions, tags storage"],
        ["Auth", "Supabase Auth", "JWT + HTTP-only cookies", "User registration, login, sessions"],
    ],
    col_widths=[1.0, 1.5, 1.8, 2.2]
)

doc.add_heading("8.2 Request Flow (Example: Create Snippet)", level=2)

flow_lines = [
    "User fills form → clicks 'Save'",
    "  │",
    "  ▼ fetch('/api/snippets', { method: 'POST' })",
    "  │",
    "  ▼ Middleware: updateSession(request)",
    "    → Refreshes auth cookies if expired",
    "  │",
    "  ▼ API Route: POST /api/snippets",
    "    1. createClient() — server Supabase client",
    "    2. supabase.auth.getUser() — verify JWT",
    "    3. INSERT INTO snippets → returns row",
    "    4. INSERT INTO snippet_versions (v1)",
    "    5. INSERT INTO snippet_tags × N",
    "    6. Return { snippet } as JSON",
    "  │",
    "  ▼ Client: SWR mutate() → re-fetch list",
    "  │",
    "  ▼ Grid re-renders with new card",
]
add_code_block(flow_lines)

add_divider()


# ── 9. TECHNOLOGY STACK ─────────────────────────────────
doc.add_heading("9. Technology Stack", level=1)

p = doc.add_paragraph()
run = p.add_run("💵 Total Cost: $0")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = GREEN
p2 = doc.add_paragraph(
    "Every technology used in this prototype is either open-source or available on a permanently free tier. "
    "No credit card required. No trial period. Completely free."
)
p2.runs[0].font.color.rgb = GRAY_TEXT

doc.add_heading("9.1 Frontend (All Open-Source)", level=2)
add_styled_table(
    ["Technology", "Version", "Purpose", "Cost"],
    [
        ["Next.js", "16.2", "React framework, App Router, SSR", "🟢 Free (MIT)"],
        ["React", "19", "UI library with hooks", "🟢 Free (MIT)"],
        ["TypeScript", "5.7", "Static typing", "🟢 Free (Apache 2.0)"],
        ["Tailwind CSS", "4.2", "Utility-first styling", "🟢 Free (MIT)"],
        ["shadcn/ui", "Latest", "Radix UI components", "🟢 Free (MIT)"],
        ["SWR", "2.2", "Data fetching + caching", "🟢 Free (MIT)"],
        ["prism-react-renderer", "2.4", "Syntax highlighting", "🟢 Free (MIT)"],
        ["diff", "7.0", "Version diff computation", "🟢 Free (BSD)"],
        ["react-hook-form", "7.54", "Form management", "🟢 Free (MIT)"],
        ["zod", "3.24", "Schema validation", "🟢 Free (MIT)"],
        ["date-fns", "4.1", "Date formatting", "🟢 Free (MIT)"],
        ["next-themes", "0.4", "Dark/light mode", "🟢 Free (MIT)"],
        ["Lucide React", "0.564", "Icon library", "🟢 Free (ISC)"],
        ["sonner", "1.7", "Toast notifications", "🟢 Free (MIT)"],
    ],
    col_widths=[1.5, 0.7, 2.3, 1.5]
)

doc.add_heading("9.2 Backend & Infrastructure (All Free Tier)", level=2)
add_styled_table(
    ["Technology", "Purpose", "Cost", "Free Tier Limits"],
    [
        ["Supabase", "PostgreSQL DB + Auth + API", "🟢 Free Plan", "500 MB DB, 50K users"],
        ["Supabase Auth", "Email/password auth + JWT", "🟢 Included", "50,000 MAU"],
        ["@supabase/ssr", "Server-side Supabase client", "🟢 Free (npm)", "—"],
        ["Vercel", "Hosting & deployment", "🟢 Hobby (Free)", "100 GB bandwidth"],
        ["Vercel Analytics", "Performance monitoring", "🟢 Included", "2,500 events/mo"],
    ],
    col_widths=[1.5, 2.0, 1.5, 1.5]
)

add_divider()


# ── 10. DATA MODEL ──────────────────────────────────────
doc.add_heading("10. Data Model & Schema Design", level=1)

doc.add_heading("10.1 Database Tables", level=2)

add_styled_table(
    ["Table", "Key Columns", "Relationships"],
    [
        ["snippets", "id, user_id, title, description, language, current_content, created_at, updated_at", "FK → auth.users (user_id)"],
        ["snippet_versions", "id, snippet_id, version_number, content, commit_message, created_at", "FK → snippets (snippet_id)"],
        ["tags", "id, user_id, name", "FK → auth.users (user_id)"],
        ["snippet_tags", "snippet_id, tag_id (composite PK)", "FK → snippets, FK → tags"],
    ],
    col_widths=[1.5, 3.5, 1.8]
)

doc.add_heading("10.2 Key Design Decisions", level=2)
add_styled_table(
    ["Decision", "Rationale"],
    [
        ["current_content in snippets", "Avoids a JOIN to snippet_versions for viewing latest code"],
        ["Sequential version_number", "Simpler than Git SHA — users understand 'v1, v2, v3'"],
        ["User-scoped tags", "Each user has their own tag namespace — no collisions"],
        ["Junction table for tags", "Allows many-to-many without JSON arrays (better querying)"],
        ["UUIDs for primary keys", "Globally unique, no sequential ID leaks"],
    ],
    col_widths=[2.5, 4.0]
)

add_divider()


# ── 11. API DESIGN ──────────────────────────────────────
doc.add_heading("11. API Design", level=1)

doc.add_heading("11.1 Endpoint Reference", level=2)

add_styled_table(
    ["Method", "Endpoint", "Purpose", "Auth Required"],
    [
        ["GET", "/api/snippets", "List snippets (with search, language, tag filters)", "✅ Yes"],
        ["POST", "/api/snippets", "Create snippet + initial version", "✅ Yes"],
        ["GET", "/api/snippets/[id]", "Get single snippet with tags", "✅ Yes"],
        ["PUT", "/api/snippets/[id]", "Update snippet, create version if content changed", "✅ Yes"],
        ["DELETE", "/api/snippets/[id]", "Delete snippet + cascade versions/tags", "✅ Yes"],
        ["GET", "/api/snippets/[id]/versions", "Get version history for a snippet", "✅ Yes"],
        ["GET", "/api/tags", "List all user tags", "✅ Yes"],
        ["POST", "/api/tags", "Create a new tag", "✅ Yes"],
    ],
    col_widths=[0.8, 2.2, 2.5, 1.0]
)

doc.add_heading("11.2 Authentication", level=2)
add_body(
    "All API routes require a valid Supabase session cookie. The middleware refreshes the session "
    "on every request. Each handler calls supabase.auth.getUser() to verify the JWT and extract the user.id. "
    "Unauthorized requests receive a 401 error response."
)

add_divider()


# ── 12. SECURITY ────────────────────────────────────────
doc.add_heading("12. Security & Access Control", level=1)

add_styled_table(
    ["Layer", "Mechanism"],
    [
        ["Authentication", "Supabase Auth with email verification, JWT tokens in HTTP-only cookies"],
        ["Session Management", "Next.js middleware refreshes session on every request"],
        ["API Authorization", "Every route handler checks supabase.auth.getUser() before processing"],
        ["Data Isolation", "All DB queries include .eq('user_id', user.id) — users see only their own data"],
        ["CSRF Protection", "SameSite cookie attributes (handled by Supabase SSR)"],
        ["Input Sanitization", "Supabase parameterized queries prevent SQL injection"],
        ["RLS (Recommended)", "Supabase Row Level Security as defense-in-depth"],
    ],
    col_widths=[2.0, 4.5]
)

add_divider()


# ── 13. TIMELINE ────────────────────────────────────────
doc.add_heading("13. Development Process & Timeline", level=1)

add_styled_table(
    ["Phase", "Duration", "Activities"],
    [
        ["Phase 1: Research & Planning", "1 week", "Market research, PRD writing, wireframes, tech stack selection"],
        ["Phase 2: Database & Auth", "1 week", "Supabase setup, schema design, migration scripts, auth flow"],
        ["Phase 3: Core API", "1 week", "CRUD endpoints for snippets, tags, versions"],
        ["Phase 4: Frontend — Dashboard", "2 weeks", "Snippet grid, search, filters, create/edit dialog, detail panel"],
        ["Phase 5: Advanced Features", "1 week", "Version history, diff viewer, syntax highlighting"],
        ["Phase 6: Polish & Testing", "1 week", "UI refinement, error handling, responsive design"],
        ["Phase 7: Deployment", "2 days", "Vercel deployment, environment variables, domain setup"],
    ],
    col_widths=[2.0, 1.0, 3.5]
)

doc.add_heading("Development Methodology", level=2)
add_bullet("Agile sprints (1 week each)")
add_bullet("Git Flow branching strategy (main → develop → feature branches)")
add_bullet("Pull request reviews before merging")
add_bullet("Continuous deployment via Vercel (auto-deploy on push to main)")

add_divider()


# ── 14. FUTURE ROADMAP ──────────────────────────────────
doc.add_heading("14. Future Roadmap (Vision Only — Not Current Scope)", level=1)

add_note_box(
    "Everything below is aspirational — documented to show product vision and growth potential for "
    "academic evaluation. None of these are planned for the current prototype. Many would require "
    "paid services, significant engineering time, or both.",
    "Disclaimer"
)

doc.add_heading("Phase 2: Intelligence Layer 🧠", level=2)
p = doc.add_paragraph("Would require paid APIs")
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = GRAY_TEXT

add_styled_table(
    ["Feature", "Description", "Estimated Cost"],
    [
        ["AI Auto-Tagging", "Analyze code on save → suggest tags", "💰 ~$5–20/month (OpenAI API)"],
        ["Semantic Search", "Natural language queries across snippets", "💰 ~$5–10/month (embeddings)"],
        ["Smart Suggestions", "'You might also need...' recommendations", "⏱️ Heavy dev time"],
        ["Code Explanation", "AI-generated description of what code does", "💰 ~$10–30/month (GPT-4o)"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_heading("Phase 3: Collaboration 🤝", level=2)
p = doc.add_paragraph("Would require Supabase Pro plan ($25/mo)")
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = GRAY_TEXT

add_styled_table(
    ["Feature", "Description", "Estimated Cost"],
    [
        ["Team Workspaces", "Shared snippet collections", "💰 Supabase Pro ($25/mo)"],
        ["Real-Time Co-Editing", "Live cursor and editing", "💰⏱️ Heavy dev + Realtime"],
        ["Comment Threads", "Discussions on snippets/versions", "⏱️ Moderate dev time"],
        ["Snippet Marketplace", "Public library with upvotes", "💰⏱️ Infrastructure + dev"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_heading("Phase 4: Ecosystem 🌐", level=2)
p = doc.add_paragraph("Separate projects entirely")
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = GRAY_TEXT

add_styled_table(
    ["Feature", "Description", "Estimated Effort"],
    [
        ["VS Code Extension", "Save/search snippets from IDE", "⏱️ 3–4 weeks"],
        ["Chrome Extension", "Right-click → save code from web", "⏱️ 2–3 weeks"],
        ["CLI Tool", "sv save, sv search, sv copy", "⏱️ 1–2 weeks"],
        ["Public API", "REST API for third-party integrations", "⏱️ 1 week"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

p = doc.add_paragraph()
run = p.add_run("Current status: ")
run.bold = True
p.add_run("Phase 1 is complete and fully operational on free infrastructure. "
          "Phases 2–4 represent where the product could evolve if it transitions from a class prototype to a real product.")

add_divider()


# ── 15. SUCCESS METRICS ─────────────────────────────────
doc.add_heading("15. Success Metrics", level=1)

doc.add_heading("Product KPIs", level=2)
add_styled_table(
    ["Metric", "Target (MVP)", "Target (6 months)"],
    [
        ["Registered Users", "50", "1,000"],
        ["Daily Active Users", "10", "200"],
        ["Snippets Created", "500", "20,000"],
        ["Average Snippets per User", "10", "20"],
        ["Average Versions per Snippet", "2", "4"],
        ["Search Queries per Day", "50", "2,000"],
        ["Copy-to-Clipboard per Day", "30", "1,500"],
    ],
    col_widths=[2.5, 1.5, 1.5]
)

doc.add_heading("Technical KPIs", level=2)
add_styled_table(
    ["Metric", "Target"],
    [
        ["Page Load Time (LCP)", "< 2.5 seconds"],
        ["API Response Time (p95)", "< 500ms"],
        ["Uptime", "99.5%"],
        ["Lighthouse Score", "> 90 (all categories)"],
        ["Error Rate", "< 1% of API requests"],
    ],
    col_widths=[3.0, 3.0]
)

add_divider()


# ── APPENDICES ──────────────────────────────────────────
doc.add_heading("Appendix A: Supported Languages", level=1)
add_body(
    "JavaScript, TypeScript, Python, Rust, Go, CSS, HTML, SQL, Bash, Java, C, C++, C#, Ruby, PHP, "
    "Swift, Kotlin, Scala, Haskell, Lua, Perl, R, Assembly, JSON, YAML, Markdown, Plain Text"
)

doc.add_heading("Appendix B: Glossary", level=1)
add_styled_table(
    ["Term", "Definition"],
    [
        ["Snippet", "A reusable piece of code stored in SnippetVault"],
        ["Version", "A historical snapshot of a snippet's code content"],
        ["Commit Message", "A description of what changed in a version"],
        ["Diff", "A visual comparison of two versions showing additions and deletions"],
        ["Tag", "A user-defined label for categorizing snippets"],
        ["SWR", "Stale-While-Revalidate — shows cached data while refreshing in background"],
        ["RLS", "Row Level Security — database-level access control policies"],
        ["SSR", "Server-Side Rendering — HTML generated on the server"],
    ],
    col_widths=[1.5, 5.0]
)

# ── END ─────────────────────────────────────────────────
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run("— End of Document —")
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = GRAY_TEXT

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run("SnippetVault PRD v1.0 • April 2026")
run.font.size = Pt(9)
run.font.color.rgb = GRAY_TEXT


# ── SAVE ────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs", "SnippetVault_PRD.docx")
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024:.1f} KB")
